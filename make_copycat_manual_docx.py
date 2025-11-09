from docx import Document

doc = Document()

# Title Page
doc.add_heading('CopyCat — AI Text Scanner Manual', level=0)
doc.add_paragraph('Operational handbook, deployment notes, reviewer guidance, demonstrations, and API reference.\n')
doc.add_paragraph('Version v0.2.x')
doc.add_paragraph('Author: Calypso Labs')
doc.add_page_break()

# Table of Contents placeholder
doc.add_heading('Table of Contents', level=1)
toc_entries = [
 "1. Overview",
 "2. UI Anatomy",
 "3. Modes & Abstain",
 "4. Runtime Settings Reference",
 "5. Demonstrations (Confidence Movers)",
 "6. Interpreting Results",
 "7. Guards & Caps (Detailed)",
 "8. Troubleshooting",
 "9. API Reference",
 "10. Deployment & Ops",
 "11. Policy & Ethics",
 "12. FAQ",
 "13. Demo Checklist"
]
for entry in toc_entries:
    doc.add_paragraph(entry, style='List Bullet')
doc.add_page_break()

def add_section(title, content):
    doc.add_heading(title, level=1)
    for paragraph in content:
        doc.add_paragraph(paragraph)

# Sections (same content)
add_section("1. Overview",[
"CopyCat estimates AI-generation likelihood using:",
"- Token-level predictability (Top-K ranks)",
"- Perplexity",
"- Burstiness",
"- Stylometry",
"- Public-domain overlap",
"- Calibrated uncertainty bands",
"Outputs include raw metrics, calibrated verdict, and abstain behavior."
])

add_section("2. UI Anatomy",[
"Left panel: text input + raw trace",
"Right panel: runtime controls (mode, caps, language gates, PD controls)",
"Mode changes calibration — NOT raw metrics."
])

add_section("3. Modes & Abstain",[
"Balanced: default sensitivity",
"Strict: higher sensitivity, narrower abstain band",
"Academic: conservative, wider abstain, safer for appeals"
])

add_section("4. Runtime Settings Reference",[
"Key controls:",
"- MODE",
"- MIN_TOKENS_STRONG",
"- SHORT_CAP / MAX_CONF_SHORT",
"- MAX_CONF_UNSTABLE",
"- EN_THRESH / NON_EN_CAP",
"- ABSTAIN_LOW / ABSTAIN_HIGH",
"- PD fingerprint settings"
])

add_section("5. Demonstrations (Confidence Movers)",[
"Test cases to illustrate behavior:",
"- Strict vs Academic on same text",
"- PD overlap cap",
"- Short-text cap behavior",
"- Nonsense guard (rhyme, nonce words)",
"- Instability cap for style-stitched samples"
])

add_section("6. Interpreting Results",[
"Likely human → proceed normally",
"Inconclusive → gather more text",
"Likely AI → contextual due process required"
])

add_section("7. Guards & Caps (Detailed)",[
"Classic style guard",
"Nonsense/verse guard",
"Public-domain cap",
"Short-text & instability caps",
"Language caps"
])

add_section("8. Troubleshooting",[
"If settings don’t change output: press Save then Rescan",
"500 error on upload → `pip install python-multipart`",
"PD overlap always 0 → fingerprints missing",
"Secondary model not blending → vocab mismatch"
])

add_section("9. API Reference",[
"POST /scan — body: {text, tag, mode}",
"GET /version",
"GET /demo",
"curl example:",
'curl -s http://127.0.0.1:8080/scan -H "content-type: application/json" -d \'{"text":"Hello world","mode":"Strict"}\''
])

add_section("10. Deployment & Ops",[
"Critical env vars include MODE, PD settings, caps, ensemble flags.",
"Fingerprint format: {'name':'corpus','ngrams':{...}}"
])

add_section("11. Policy & Ethics",[
"CopyCat does not prove authorship.",
"Use as advisory evidence with fair process."
])

add_section("12. FAQ",[
"Q: Does mode change PPL? A: No — only calibration.",
"Q: Can CopyCat prove authorship? A: No.",
"Q: Why capped at 0.25? A: PD or short-text cap likely.",
])

add_section("13. Demo Checklist",[
"Classic-limiter test",
"Nonsense guard",
"Short-text cap",
"PD overlap",
"Strict vs Academic spread"
])

doc.save("CopyCat_Manual_v0_2x.docx")
print("✅ Saved file: CopyCat_Manual_v0_2x.docx")

